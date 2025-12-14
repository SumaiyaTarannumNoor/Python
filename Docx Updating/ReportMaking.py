from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Training data from your output
training_data = """Epoch 001 | Train MSE=1.05034 | Val MSE=0.82228 | Val Acc=17.95%
Epoch 002 | Train MSE=0.95591 | Val MSE=0.74966 | Val Acc=25.20%
Epoch 003 | Train MSE=0.87015 | Val MSE=0.68373 | Val Acc=31.77%
Epoch 004 | Train MSE=0.79227 | Val MSE=0.62388 | Val Acc=37.75%
Epoch 005 | Train MSE=0.72154 | Val MSE=0.56955 | Val Acc=43.17%
Epoch 006 | Train MSE=0.65730 | Val MSE=0.52023 | Val Acc=48.09%
Epoch 007 | Train MSE=0.59896 | Val MSE=0.47546 | Val Acc=52.56%
Epoch 008 | Train MSE=0.54598 | Val MSE=0.43482 | Val Acc=56.61%
Epoch 009 | Train MSE=0.49787 | Val MSE=0.39793 | Val Acc=60.29%
Epoch 010 | Train MSE=0.45417 | Val MSE=0.36444 | Val Acc=63.63%
Epoch 011 | Train MSE=0.41448 | Val MSE=0.33405 | Val Acc=66.67%
Epoch 012 | Train MSE=0.37844 | Val MSE=0.30646 | Val Acc=69.42%
Epoch 013 | Train MSE=0.34570 | Val MSE=0.28142 | Val Acc=71.92%
Epoch 014 | Train MSE=0.31598 | Val MSE=0.25869 | Val Acc=74.19%
Epoch 015 | Train MSE=0.28898 | Val MSE=0.23807 | Val Acc=76.24%
Epoch 016 | Train MSE=0.26445 | Val MSE=0.21935 | Val Acc=78.11%
Epoch 017 | Train MSE=0.24218 | Val MSE=0.20236 | Val Acc=79.81%
Epoch 018 | Train MSE=0.22195 | Val MSE=0.18694 | Val Acc=81.35%
Epoch 019 | Train MSE=0.20358 | Val MSE=0.17295 | Val Acc=82.74%
Epoch 020 | Train MSE=0.18690 | Val MSE=0.16025 | Val Acc=84.01%
Epoch 021 | Train MSE=0.17174 | Val MSE=0.14873 | Val Acc=85.16%
Epoch 022 | Train MSE=0.15798 | Val MSE=0.13827 | Val Acc=86.20%
Epoch 023 | Train MSE=0.14548 | Val MSE=0.12879 | Val Acc=87.15%
Epoch 024 | Train MSE=0.13412 | Val MSE=0.12018 | Val Acc=88.01%
Epoch 025 | Train MSE=0.12381 | Val MSE=0.11237 | Val Acc=88.79%
Epoch 026 | Train MSE=0.11444 | Val MSE=0.10528 | Val Acc=89.49%
Epoch 027 | Train MSE=0.10593 | Val MSE=0.09885 | Val Acc=90.14%
Epoch 028 | Train MSE=0.09820 | Val MSE=0.09302 | Val Acc=90.72%
Epoch 029 | Train MSE=0.09118 | Val MSE=0.08773 | Val Acc=91.25%
Epoch 030 | Train MSE=0.08480 | Val MSE=0.08293 | Val Acc=91.73%
Epoch 031 | Train MSE=0.07900 | Val MSE=0.07857 | Val Acc=92.16%
Epoch 032 | Train MSE=0.07374 | Val MSE=0.07462 | Val Acc=92.55%
Epoch 033 | Train MSE=0.06895 | Val MSE=0.07104 | Val Acc=92.91%
Epoch 034 | Train MSE=0.06461 | Val MSE=0.06779 | Val Acc=93.24%
Epoch 035 | Train MSE=0.06066 | Val MSE=0.06484 | Val Acc=93.53%
Epoch 036 | Train MSE=0.05707 | Val MSE=0.06217 | Val Acc=93.80%
Epoch 037 | Train MSE=0.05382 | Val MSE=0.05975 | Val Acc=94.04%
Epoch 038 | Train MSE=0.05085 | Val MSE=0.05755 | Val Acc=94.26%
Epoch 039 | Train MSE=0.04816 | Val MSE=0.05555 | Val Acc=94.46%
Epoch 040 | Train MSE=0.04572 | Val MSE=0.05374 | Val Acc=94.64%
Epoch 041 | Train MSE=0.04349 | Val MSE=0.05210 | Val Acc=94.80%
Epoch 042 | Train MSE=0.04147 | Val MSE=0.05062 | Val Acc=94.95%
Epoch 043 | Train MSE=0.03963 | Val MSE=0.04927 | Val Acc=95.08%
Epoch 044 | Train MSE=0.03796 | Val MSE=0.04805 | Val Acc=95.21%
Epoch 045 | Train MSE=0.03645 | Val MSE=0.04694 | Val Acc=95.32%
Epoch 046 | Train MSE=0.03506 | Val MSE=0.04593 | Val Acc=95.42%
Epoch 047 | Train MSE=0.03381 | Val MSE=0.04502 | Val Acc=95.51%
Epoch 048 | Train MSE=0.03267 | Val MSE=0.04419 | Val Acc=95.59%
Epoch 049 | Train MSE=0.03163 | Val MSE=0.04345 | Val Acc=95.66%
Epoch 050 | Train MSE=0.03068 | Val MSE=0.04277 | Val Acc=95.73%
Epoch 051 | Train MSE=0.02982 | Val MSE=0.04215 | Val Acc=95.79%
Epoch 052 | Train MSE=0.02904 | Val MSE=0.04159 | Val Acc=95.85%
Epoch 053 | Train MSE=0.02833 | Val MSE=0.04108 | Val Acc=95.90%
Epoch 054 | Train MSE=0.02768 | Val MSE=0.04062 | Val Acc=95.95%
Epoch 055 | Train MSE=0.02709 | Val MSE=0.04021 | Val Acc=95.99%
Epoch 056 | Train MSE=0.02655 | Val MSE=0.03983 | Val Acc=96.03%
Epoch 057 | Train MSE=0.02606 | Val MSE=0.03948 | Val Acc=96.06%
Epoch 058 | Train MSE=0.02561 | Val MSE=0.03917 | Val Acc=96.09%
Epoch 059 | Train MSE=0.02520 | Val MSE=0.03889 | Val Acc=96.12%
Epoch 060 | Train MSE=0.02483 | Val MSE=0.03863 | Val Acc=96.15%
Epoch 061 | Train MSE=0.02449 | Val MSE=0.03840 | Val Acc=96.17%
Epoch 062 | Train MSE=0.02418 | Val MSE=0.03819 | Val Acc=96.19%
Epoch 063 | Train MSE=0.02389 | Val MSE=0.03799 | Val Acc=96.21%
Epoch 064 | Train MSE=0.02363 | Val MSE=0.03782 | Val Acc=96.23%
Epoch 065 | Train MSE=0.02340 | Val MSE=0.03766 | Val Acc=96.24%
Epoch 066 | Train MSE=0.02318 | Val MSE=0.03751 | Val Acc=96.26%
Epoch 067 | Train MSE=0.02298 | Val MSE=0.03738 | Val Acc=96.27%
Epoch 068 | Train MSE=0.02280 | Val MSE=0.03726 | Val Acc=96.28%
Epoch 069 | Train MSE=0.02263 | Val MSE=0.03715 | Val Acc=96.29%
Epoch 070 | Train MSE=0.02248 | Val MSE=0.03705 | Val Acc=96.30%
Epoch 071 | Train MSE=0.02234 | Val MSE=0.03696 | Val Acc=96.31%
Epoch 072 | Train MSE=0.02221 | Val MSE=0.03687 | Val Acc=96.32%
Epoch 073 | Train MSE=0.02209 | Val MSE=0.03680 | Val Acc=96.33%
Epoch 074 | Train MSE=0.02198 | Val MSE=0.03673 | Val Acc=96.34%
Epoch 075 | Train MSE=0.02187 | Val MSE=0.03666 | Val Acc=96.34%
Epoch 076 | Train MSE=0.02178 | Val MSE=0.03660 | Val Acc=96.35%
Epoch 077 | Train MSE=0.02169 | Val MSE=0.03655 | Val Acc=96.35%
Epoch 078 | Train MSE=0.02161 | Val MSE=0.03650 | Val Acc=96.36%
Epoch 079 | Train MSE=0.02154 | Val MSE=0.03645 | Val Acc=96.36%
Epoch 080 | Train MSE=0.02147 | Val MSE=0.03641 | Val Acc=96.37%"""

# Parse the training data
def parse_training_data(data):
    lines = data.strip().split('\n')
    parsed = []
    for line in lines:
        parts = line.split('|')
        epoch = parts[0].strip().split()[-1]
        train_mse = parts[1].strip().split('=')[1]
        val_mse = parts[2].strip().split('=')[1]
        val_acc = parts[3].strip().split('=')[1].replace('%', '')
        parsed.append((epoch, train_mse, val_mse, val_acc))
    return parsed

# Load the existing document
doc = Document('Project_02_Electricity_Load_Forecasting_Report.docx')

# Find and replace the table
table_found = False
for i, table in enumerate(doc.tables):
    # Check if this is the training metrics table (it should have 4 columns)
    if len(table.columns) == 4:
        # Check if first row contains headers like "Epoch", "Train MSE", etc.
        header_row = table.rows[0]
        if 'Epoch' in header_row.cells[0].text or 'Train MSE' in header_row.cells[1].text:
            table_found = True
            
            # Clear existing data rows (keep header)
            for _ in range(len(table.rows) - 1):
                table._element.remove(table.rows[-1]._element)
            
            # Parse and add new data
            data = parse_training_data(training_data)
            
            for epoch, train_mse, val_mse, val_acc in data:
                row_cells = table.add_row().cells
                row_cells[0].text = epoch
                row_cells[1].text = train_mse
                row_cells[2].text = val_mse
                row_cells[3].text = val_acc
                
                # Center align all cells
                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            print(f"✓ Updated table with {len(data)} epochs of training data")
            break

if not table_found:
    print("Warning: Could not find the training metrics table to update")

# Update the best validation accuracy and test accuracy text
for paragraph in doc.paragraphs:
    if "Best Validation Accuracy:" in paragraph.text:
        paragraph.text = "Best Validation Accuracy: 96.37% (Epoch 80)"
    if "Final Test Accuracy:" in paragraph.text:
        paragraph.text = "Final Test Accuracy: 99.33%"

# Save the updated document
output_filename = 'Project_02_Electricity_Load_Forecasting_Report_Updated.docx'
doc.save(output_filename)

print(f"\n✓ Document saved as '{output_filename}'")
print(f"✓ Updated with 80 epochs of training data")
print(f"✓ Best Validation Accuracy: 96.37%")
print(f"✓ Final Test Accuracy: 99.33%")